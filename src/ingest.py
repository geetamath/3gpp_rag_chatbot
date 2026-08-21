"""
Step 1 + 2: Document Ingestion and Clause-Aware Chunking.
"""

import os
import re
import json
import hashlib
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import config

CLAUSE_RE = re.compile(r"^(?P<num>\d+(?:\.\d+){0,5})\s+(?P<title>[A-Z][^\n]{2,120})$")
TS_RE = re.compile(r"(TS|TR)\s?(\d{2}\.\d{3})")
TS_FILENAME_RE = re.compile(r"(?<!\d)(\d{2})(\d{3})[-_]")
RELEASE_PAREN_RE = re.compile(r"\(Release\s+(\d{1,2})\)", re.IGNORECASE)
RELEASE_RE = re.compile(r"Rel(?:ease)?[-\s]?(\d{1,2})", re.IGNORECASE)
_TITLE_BOILERPLATE_RE = re.compile(
    r"^\s*(3GPP|3rd Generation|Technical Specification|Stage\s+\d|"
    r"TS\s?\d|TR\s?\d|V\d+\.\d+\.\d+|\(Release|\d{1,4}\s*$)",
    re.IGNORECASE,
)


def _approx_tokens(text: str) -> int:
    return max(1, len(text.split()))


def _infer_ts_and_release(filename: str, text_head: str, header_text: str = ""):
    ts_number, release = None, None
    m = TS_RE.search(filename) or TS_RE.search(header_text) or TS_RE.search(text_head)
    if m:
        ts_number = f"{m.group(1)} {m.group(2)}"
    if not ts_number:
        m3 = TS_FILENAME_RE.search(filename)
        if m3:
            ts_number = f"TS {m3.group(1)}.{m3.group(2)}"

    m2 = (
        RELEASE_PAREN_RE.search(header_text)
        or RELEASE_RE.search(header_text)
        or RELEASE_PAREN_RE.search(text_head)
        or RELEASE_RE.search(filename)
    )
    if m2:
        release = f"Rel-{m2.group(1)}"
    return ts_number or "UNKNOWN", release or "UNKNOWN"


def _infer_doc_title(text_head: str) -> str:
    m = RELEASE_PAREN_RE.search(text_head)
    if not m:
        return ""
    window_start = max(0, m.start() - 400)
    window_end = min(len(text_head), m.end() + 400)
    before = text_head[window_start:m.start()]
    after = text_head[m.end():window_end]
    before_lines = [ln.strip() for ln in before.splitlines() if ln.strip()]
    after_lines = [ln.strip() for ln in after.splitlines() if ln.strip()]
    candidates = []
    for i in range(max(len(before_lines), len(after_lines))):
        if i < len(before_lines):
            candidates.append(before_lines[-(i + 1)])
        if i < len(after_lines):
            candidates.append(after_lines[i])
    for line in candidates:
        if _TITLE_BOILERPLATE_RE.match(line):
            continue
        if len(line) < 8:
            continue
        return line.rstrip(";").strip()[:200]
    return ""


def _read_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _read_docx_header(path: str) -> str:
    """Read only the page header text (not body/ToC) -- this is where 3GPP
    puts the reliable 'Release N' + 'V<major>.x.x' line, avoiding false
    matches on body/ToC lines like 'NAS signalling connection Release  111'
    (a heading literally ending in the word 'Release', followed by a page
    number that a loose body-text regex can mistake for a release number).
    """
    from docx import Document
    doc = Document(path)
    lines = []
    for section in doc.sections:
        for para in section.header.paragraphs:
            t = para.text.strip()
            if t:
                lines.append(t)
    return "\n".join(lines)


def _read_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    lines = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        for line in text.split("\n"):
            line = line.strip()
            if line:
                lines.append(f"{line}\x00PAGE{i+1}")
    return "\n".join(lines)


def _read_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_raw_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        return _read_docx(path)
    if ext == ".pdf":
        return _read_pdf(path)
    if ext in (".txt", ".md"):
        return _read_txt(path)
    raise ValueError(f"Unsupported file type: {ext}")


def split_into_clauses(raw_text: str):
    clauses = []
    current = {"num": "0", "title": "Preamble", "lines": [], "page": None}

    for raw_line in raw_text.split("\n"):
        page = None
        line = raw_line
        if "\x00PAGE" in raw_line:
            line, page_marker = raw_line.split("\x00PAGE")
            page = int(page_marker)

        m = CLAUSE_RE.match(line.strip())
        if m:
            if current["lines"]:
                clauses.append(current)
            current = {"num": m.group("num"), "title": m.group("title").strip(),
                       "lines": [], "page": page}
        else:
            if line.strip():
                current["lines"].append(line.strip())
                if page and current["page"] is None:
                    current["page"] = page

    if current["lines"]:
        clauses.append(current)

    return [(c["num"], c["title"], "\n".join(c["lines"]), c["page"]) for c in clauses]


def chunk_clause(clause_id, clause_title, body_text, page):
    breadcrumb = f"{clause_id} {clause_title}".strip()
    n_tokens = _approx_tokens(body_text)

    if n_tokens <= config.MAX_CHUNK_TOKENS:
        return [{"clause_id": clause_id, "clause_title": clause_title,
                  "breadcrumb": breadcrumb, "text": body_text, "page": page}]

    words = body_text.split()
    step = config.MAX_CHUNK_TOKENS - config.CHUNK_OVERLAP_TOKENS
    out = []
    for start in range(0, len(words), step):
        piece = " ".join(words[start:start + config.MAX_CHUNK_TOKENS])
        if not piece.strip():
            continue
        out.append({"clause_id": clause_id, "clause_title": clause_title,
                     "breadcrumb": breadcrumb, "text": piece, "page": page})
    return out


def ingest_file(path: str):
    filename = os.path.basename(path)
    raw_text = load_raw_text(path)

    header_text = ""
    if os.path.splitext(filename)[1].lower() == ".docx":
        try:
            header_text = _read_docx_header(path)
        except Exception as e:
            print(f"[ingest] WARNING: could not read header for {filename}: {e}", flush=True)

    ts_number, release = _infer_ts_and_release(filename, raw_text[:20000], header_text)
    doc_title = _infer_doc_title(raw_text[:20000])

    clauses = split_into_clauses(raw_text)

    chunks = []
    for clause_id, clause_title, body_text, page in clauses:
        if _approx_tokens(body_text) < 5:
            continue
        for sub in chunk_clause(clause_id, clause_title, body_text, page):
            uid_src = f"{filename}|{sub['clause_id']}|{sub['text'][:50]}"
            chunk_id = hashlib.sha1(uid_src.encode()).hexdigest()[:16]
            chunks.append({
                "chunk_id": chunk_id,
                "ts_number": ts_number,
                "release": release,
                "doc_title": doc_title,
                "clause_id": sub["clause_id"],
                "clause_title": sub["clause_title"],
                "breadcrumb": sub["breadcrumb"],
                "page": sub["page"],
                "text": sub["text"],
                "source_file": filename,
            })
    return chunks


def ingest_directory(directory: str):
    all_chunks = []
    if not os.path.isdir(directory):
        return all_chunks
    for fn in sorted(os.listdir(directory)):
        path = os.path.join(directory, fn)
        if not os.path.isfile(path):
            continue
        if os.path.splitext(fn)[1].lower() not in (".docx", ".pdf", ".txt", ".md"):
            continue
        try:
            all_chunks.extend(ingest_file(path))
        except Exception as e:
            print(f"[ingest] WARNING: failed on {fn}: {e}")
    return all_chunks


def save_chunks(chunks, out_path=config.CHUNKS_PATH):
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        for c in chunks:
            f.write(json.dumps(c, ensure_ascii=False) + "\n")
    print(f"[ingest] wrote {len(chunks)} chunks -> {out_path}")


def load_chunks(path=config.CHUNKS_PATH):
    chunks = []
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if line:
                chunks.append(json.loads(line))
    return chunks


if __name__ == "__main__":
    source_dir = (
        config.RAW_DIR
        if os.path.isdir(config.RAW_DIR) and os.listdir(config.RAW_DIR)
        else config.SAMPLE_DIR
    )
    print(f"[ingest] ingesting from: {source_dir}")
    chunks = ingest_directory(source_dir)
    save_chunks(chunks)